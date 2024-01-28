"""
****************************************************************************************************
    * @file	    :   SignalViewer.py
    * @brief	:   A Desktop Application to view signals and generate a report of them for analysis
    * @authors	:   Mohamed Sami Ahmed
                    Mohamed Sayed Abd El-Salam
                    Kareem Salah Noureddine
****************************************************************************************************
"""
import os
import random
import sys
import webbrowser
from io import BytesIO

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import pyqtgraph as pg
import wfdb
from docx import Document
from docx2pdf import convert
from docx.shared import Inches
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
from PyQt5 import QtGui
from PyQt5.QtCore import QPoint, Qt, QTimer
from PyQt5.QtGui import QBrush, QColor, QCursor, QIcon
from PyQt5.QtWidgets import QSplitter  # Use QSplitter to divide the UI into sections
from PyQt5.QtWidgets import (
    QAction,
    QApplication,
    QCheckBox,
    QColorDialog,
    QComboBox,
    QFileDialog,
    QHBoxLayout,
    QHeaderView,
    QLabel,
    QMainWindow,
    QMenu,
    QMessageBox,
    QPushButton,
    QSlider,
    QTableWidget,
    QTableWidgetItem,
    QVBoxLayout,
    QWidget,
)
from pyqtgraph.exporters import ImageExporter


class SignalViewer(QMainWindow):
    def __init__(self):
        super().__init__()

        # --- Main Window Initialization --- #
        # ---------------------------------- #
        self.window_width, self.window_height = 1000, 800
        self.resize(self.window_width, self.window_height)
        self.setWindowTitle("Signal Viewer")
        self.setWindowIcon(QtGui.QIcon("Resources/Icons/logo.png"))
        self.setAcceptDrops(True)

        # --- Creating a Menubar --- #
        # -------------------------- #
        self.menuBar = self.menuBar()
        file_menu = self.menuBar.addMenu("File")
        view_menu = self.menuBar.addMenu("View")
        help_menu = self.menuBar.addMenu("Help")

        import_action = QAction("Import Signal", self)
        import_action.triggered.connect(self.import_signal)
        import_action.setShortcut("Ctrl+I")

        exit_action = QAction("Exit App", self)
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(lambda: QApplication.quit())

        file_menu.addAction(import_action)
        file_menu.addAction(exit_action)

        pdf_report = QAction("Generate PDF Report", self)
        pdf_report.triggered.connect(self.convert_to_pdf)
        pdf_report.setShortcut("Ctrl+R")

        view_menu.addAction(pdf_report)

        appDocumentationAction = QAction("App Documentation", self)
        appDocumentationAction.triggered.connect(self.openDocumentation)

        help_menu.addAction(appDocumentationAction)

        # --- Creating Main Layout --- #
        # ---------------------------- #
        central_widget = QSplitter(Qt.Horizontal)
        self.setCentralWidget(central_widget)

        # Create a left widget to hold the tables and graphs
        left_widget = QWidget()
        left_widget_layout = QVBoxLayout(left_widget)
        central_widget.addWidget(left_widget)

        # Create a middle widget to hold the tables and graphs
        middle_widget = QWidget()
        middle_widget_layout = QVBoxLayout(middle_widget)
        central_widget.addWidget(middle_widget)

        # Create a right widget to hold the buttons
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        central_widget.addWidget(right_widget)

        # --- First Section (Left) - Two Tables --- #
        # ----------------------------------------- #
        # Create a layout to hold the two tables vertically aligned
        left_tables_layout = QVBoxLayout()
        left_widget_layout.addLayout(left_tables_layout)

        # Create two tables and add them to the layout
        self.signals_info_table_1 = QTableWidget(0, 4)
        self.signals_info_table_2 = QTableWidget(0, 4)

        # Create column labels
        column_labels = ["Signal", "Graph No", "Color", "Visibility"]
        self.signals_info_table_1.setHorizontalHeaderLabels(column_labels)
        self.signals_info_table_2.setHorizontalHeaderLabels(column_labels)

        # Set the item flags for each column
        for table in [self.signals_info_table_1, self.signals_info_table_2]:
            table.setSortingEnabled(False)
            table.verticalHeader().setVisible(False)
            for col in range(1, 4):  # Skip the "Signal" column (column 0)
                for row in range(0):
                    item = QTableWidgetItem("")
                    item.setFlags(item.flags() & ~Qt.ItemIsEditable)
                    table.setItem(row, col, item)

        # to ensure that columns total width fill the whole width of the table
        header_1 = self.signals_info_table_1.horizontalHeader()
        header_1.setSectionResizeMode(QHeaderView.Stretch)
        header_1.setSectionResizeMode(1, QHeaderView.ResizeToContents)

        header_2 = self.signals_info_table_2.horizontalHeader()
        header_2.setSectionResizeMode(QHeaderView.Stretch)
        header_2.setSectionResizeMode(1, QHeaderView.ResizeToContents)

        # Add the tables to the layout
        left_tables_layout.addWidget(self.signals_info_table_1)
        left_tables_layout.addWidget(self.signals_info_table_2)

        # --- Second Section (Middle) - Graphs --- #
        # --------------------------------------- #

        # Create two PlotWidgets for graphs
        self.plot_widget_1 = pg.PlotWidget()
        self.plot_widget_2 = pg.PlotWidget()
        self.plot_widget_2.setXLink(self.plot_widget_1)
        self.plot_widget_2.setYLink(self.plot_widget_1)

        # Additional graph setup for plot_widget_1
        self.plot_widget_1.setBackground("black")
        self.plot_1 = self.plot_widget_1.plot()
        self.plot_1 = self.plot_widget_1.plot()
        self.plot_widget_1.setYRange(-1, 1)
        self.plot_widget_1.setXRange(0, 1, padding=0)

        # Additional graph setup for plot_widget_2
        self.plot_widget_2.setBackground("black")
        self.plot_2 = self.plot_widget_2.plot()
        self.plot_widget_2.setYRange(-1, 1)
        self.plot_widget_2.setXRange(0, 1, padding=0)

        # Create a QTimer for each graph to update the plots
        self.timer_1 = QTimer(self)
        self.timer_1.timeout.connect(self.update_plot_1)

        self.timer_2 = QTimer(self)
        self.timer_2.timeout.connect(self.update_plot_2)

        # Create a layout to hold the two graph widgets
        middle_layout = QVBoxLayout()
        middle_layout.addWidget(self.plot_widget_1)
        middle_layout.addWidget(self.plot_widget_2)

        # Add the layout to the middle section
        middle_widget_layout.addLayout(middle_layout)

        # --- Third Section (Right) - Buttons --- #
        # --------------------------------------- #
        # Create buttons, checkboxes, and other controls
        self.pause_play_button = QPushButton(self)
        self.pause_play_button.setIcon(QIcon("Resources/Icons/pause_button.png"))
        self.pause_play_button.setShortcut("Space")
        self.pause_play_button.setCheckable(True)
        self.pause_play_button.toggled.connect(self.pause_play_toggle_event)

        self.zoom_in_button = QPushButton(self)
        self.zoom_in_button.setIcon(QIcon("Resources/Icons/zoom_in.png"))
        self.zoom_in_button.clicked.connect(self.zoom_in_event)

        self.zoom_out_button = QPushButton(self)
        self.zoom_out_button.setIcon(QIcon("Resources/Icons/zoom_out.png"))
        self.zoom_out_button.clicked.connect(self.zoom_out_event)

        self.reset_button = QPushButton()
        self.reset_button.setIcon(QIcon("Resources/Icons/reset_button.png"))
        self.reset_button.clicked.connect(self.reset_signal)

        self.linked_graphs = False
        self.link_graphs_checkbox = QCheckBox("Link Graphs")
        self.link_graphs_checkbox.setShortcut("Ctrl+L")
        self.link_graphs_checkbox.stateChanged.connect(self.link_graphs_changed)

        # Create a QLabel for "Speed: x1" and set it to be vertical
        self.speed_label = QLabel("Speed: x1")
        self.speed_label.setAlignment(Qt.AlignTop)  # Align controls to the top

        # Create a horizontal slider for speed
        self.speed_slider = QSlider(Qt.Horizontal)  # Use a horizontal slider
        self.speed_slider.setFixedWidth(
            150
        )  # Set a fixed width for the horizontal slider
        self.speed_slider.setMinimum(-50)
        self.speed_slider.setMaximum(49)
        self.speed_slider.setValue(1)
        self.speed_slider.setTickInterval(1)
        self.speed_slider.setTickPosition(QSlider.TicksBelow)
        self.speed_slider.valueChanged.connect(self.update_speed)

        self.graph_selector = QComboBox(self)
        self.graph_selector.addItems(["Graph 1", "Graph 2"])
        self.graph_selector.currentIndexChanged.connect(self.graph_selected)

        self.guidance_label = QLabel("Select which graph to affect:", self)
        self.link_graphs_guidance = QLabel("Or link graphs", self)
        self.controls = QLabel("Graph Controls:")

        # Create the "Switch Channel" button
        take_snapshot_button = QPushButton("Take Snapshot", self)
        take_snapshot_button.clicked.connect(self.take_snapshot_event)

        # Create a layout for the controls with vertical alignment
        buttons_layout = QVBoxLayout()
        buttons_layout.addWidget(self.guidance_label)
        buttons_layout.addWidget(self.graph_selector)
        buttons_layout.addSpacing(50)
        buttons_layout.addWidget(self.link_graphs_guidance)
        buttons_layout.addWidget(self.link_graphs_checkbox)
        buttons_layout.setAlignment(Qt.AlignTop)
        buttons_layout.addSpacing(50)
        buttons_layout.addWidget(self.controls)
        buttons_layout.addWidget(self.pause_play_button)
        buttons_layout.addWidget(self.reset_button)
        buttons_layout.addWidget(self.zoom_in_button)
        buttons_layout.addWidget(self.zoom_out_button)
        buttons_layout.addWidget(self.speed_label)
        buttons_layout.addWidget(self.speed_slider)
        buttons_layout.addSpacing(50)
        buttons_layout.addWidget(take_snapshot_button)

        # Add the buttons layout to the right section
        right_layout.addLayout(buttons_layout)

        # Create a dictionary to map graph selector indices to corresponding timers and plots
        self.graph_map = {
            0: {"timer": self.timer_1, "plot": self.plot_1},
            1: {"timer": self.timer_2, "plot": self.plot_2},
        }

        # Initial playing state for both graphs
        self.playing_state = {0: True, 1: True}
        # Initially, Graph 1 is selected
        self.current_graph = 0

        # Store the last position for each graph
        self.last_position = {0: 0, 1: 0}

        # List to store imported signal data and associated graph numbers
        self.imported_signals = []
        self.signal_index_1 = 0  # Initialize signal_index_1 to 0
        self.signal_index_2 = 0  # Initialize signal_index_2 to 0

        self.signal_data_1 = []  # Initialize signal_index_1 to 0
        self.signal_data_2 = []  # Initialize signal_index_2 to 0

        self.plot_items_1 = []  # For Graph 1
        self.plot_items_2 = []  # For Graph 2

        # List to store imported file names and associated graph numbers
        self.imported_files = []
        self.table_1 = []
        self.table_2 = []

        # Add a list to store the colors associated with each signal
        self.signal_colors = []

    # --- Import and Plotting Methods --- #
    # ----------------------------------- #
    def import_signal(self):
        options = QFileDialog.Options()
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "Open Signal Files",
            "",
            "Signal Files (*.csv *.hea *.dat);;All Files (*)",
            options=options,
        )

        if file_paths:
            for file_path in file_paths:
                try:
                    selected_graph = 0

                    if file_path.endswith(".hea") or file_path.endswith(".dat"):
                        # Read signal data from .hea and .dat files using wfdb library
                        record = wfdb.rdrecord(
                            file_path[:-4]
                        )  # Remove ".hea" extension
                        signal_data = record.p_signal[:, 0]  # Use the first channel

                    elif file_path.endswith(".csv"):
                        # Use pandas to read the CSV file
                        data_frame = pd.read_csv(file_path)
                        # Assuming the signal data is in a column named "Signal" in the CSV file
                        signal_data = data_frame["values"].to_numpy()

                    # Store the imported signal data and its associated graph number
                    self.imported_signals.append((signal_data, selected_graph))

                    # Ensure the self.signal_colors list has enough elements
                    self.ensure_signal_colors_length(len(self.imported_files))

                    # Choose a color for this imported signal (e.g., based on the index)
                    color = self.get_random_signal_color(len(self.imported_files))
                    self.signal_colors.append(color)

                    # Reset the selected graph's data and X-axis range
                    selected_timer = self.graph_map[selected_graph]["timer"]
                    selected_timer.stop()

                    if selected_graph == 0:
                        self.signal_data_1 = signal_data
                        self.signal_index_1 = 0  # Set the index to 0
                        self.table_1.append((file_path, selected_graph))

                    elif selected_graph == 1:
                        self.signal_data_2 = signal_data
                        self.signal_index_2 = 0  # Set the index to 0
                        self.table_2.append((file_path, selected_graph))

                    selected_timer.start(60)  # Start the timer for the selected graph

                    # Add the imported file name and associated graph number to the list
                    self.imported_files.append((file_path, selected_graph))

                except Exception as e:
                    print(f"Error loading the file: {e}")

            self.update_signal_list()

    def update_plot_1(self):
        self.plot_widget_1.clear()
        if (
            self.playing_state[0]
            and self.signal_data_1 is not None
            and self.signal_index_1 < len(self.signal_data_1)
        ):
            try:
                self.signal_index_1 += 1

                # Iterate through imported signals and plot them with different colors
                for i, (signal_data, graph_number) in enumerate(self.imported_signals):
                    if graph_number == 0:
                        x_data = np.arange(self.signal_index_1)
                        y_data = signal_data[: self.signal_index_1]

                        # Get the color for this signal from the self.signal_colors list
                        color = pg.mkColor(self.signal_colors[i])

                        # Calculate the Y-axis range based on the current data
                        y_max = max(y_data)
                        y_min = min(y_data)
                        # Check the visibility status for this signal
                        visibility_checkbox = self.signals_info_table_1.cellWidget(i, 3)
                        if visibility_checkbox.isChecked():
                            self.plot_widget_1.plot(x=x_data, y=y_data, pen=color)

                        # Set the Y-axis range for the plot
                        self.plot_widget_1.setYRange(y_min, y_max, padding=0.1)

                        # Calculate the visible range for the X-axis based on the current signal_index_1
                        visible_range = (self.signal_index_1 - 150, self.signal_index_1)
                        # Set the X-axis limits to control the visible range
                        x_min_limit = 0  # Set the minimum X-axis limit
                        x_max_limit = (
                            self.signal_index_1 + 0.1
                        )  # Set the maximum X-axis limit
                        self.plot_widget_1.setLimits(
                            xMin=x_min_limit, xMax=x_max_limit, yMin=y_min, yMax=y_max
                        )
                        self.plot_widget_1.setXRange(*visible_range, padding=0)
            except Exception as e:
                print(f"Error updating the plot for graph 1: {e}")
        pass

    def update_plot_2(self):
        self.plot_widget_2.clear()
        if (
            self.playing_state[1]
            and self.signal_data_2 is not None
            and self.signal_index_2 < len(self.signal_data_2)
        ):
            try:
                self.signal_index_2 += 1

                # Iterate through imported signals and plot them with different colors
                for i, (signal_data, graph_number) in enumerate(self.imported_signals):
                    if graph_number == 1:
                        x_data = np.arange(self.signal_index_2)
                        y_data = signal_data[: self.signal_index_2]

                        # Get the color for this signal from the self.signal_colors list
                        color = pg.mkColor(self.signal_colors[i])
                        # Calculate the Y-axis range based on the current data
                        y_max = max(y_data)
                        y_min = min(y_data)

                        # Check the visibility status for this signal
                        visibility_checkbox = self.signals_info_table_2.cellWidget(i, 3)
                        if visibility_checkbox.isChecked():
                            self.plot_widget_2.plot(x=x_data, y=y_data, pen=color)

                        # Scroll the plot to keep the most recent data in view
                        # Set the Y-axis range for the plot
                        self.plot_widget_2.setYRange(y_min, y_max, padding=0.1)

                        # Calculate the visible range for the X-axis based on the current signal_index_1
                        visible_range = (self.signal_index_2 - 150, self.signal_index_2)
                        # Set the X-axis limits to control the visible range
                        x_min_limit = 0  # Set the minimum X-axis limit
                        x_max_limit = (
                            self.signal_index_2 + 0.1
                        )  # Set the maximum X-axis limit
                        self.plot_widget_2.setLimits(
                            xMin=x_min_limit, xMax=x_max_limit, yMin=y_min, yMax=y_max
                        )
                        self.plot_widget_2.setXRange(*visible_range, padding=0)

            except Exception as e:
                print(f"Error updating the plot for graph 2: {e}")
        pass

    # --- Update The Signals Information Table --- #
    # -------------------------------------------- #

    def update_signal_list(self):
        self.signals_info_table_1.setRowCount(len(self.table_1))
        self.signals_info_table_2.setRowCount(len(self.imported_files))

        for row, (file_path, graph_number) in enumerate(self.imported_files):
            file_name = file_path.split("/")[-1]  # Extract the file name
            graph_label = QTableWidgetItem(f"Graph {graph_number + 1}")

            # Create a QTableWidgetItem and set its text
            file_name_label = QTableWidgetItem(file_name)

            # Convert the color from self.signal_colors to QColor
            color = QColor(self.signal_colors[row])

            # Create a QTableWidgetItem and set its text
            color_label = QTableWidgetItem()
            color_label.setBackground(
                QBrush(color)
            )  # Set the background color using QBrush

            if graph_number == 0:
                self.fill_table_row(
                    row,
                    self.signals_info_table_1,
                    file_name_label,
                    graph_label,
                    color_label,
                )
            elif graph_number == 1:
                self.fill_table_row(
                    row,
                    self.signals_info_table_2,
                    file_name_label,
                    graph_label,
                    color_label,
                )

    def fill_table_row(self, row, table, name, graph, color):
        table.setItem(row, 0, name)  # Display file name
        table.setItem(row, 1, graph)  # Display graph number
        table.setItem(row, 2, color)  # Display color
        # Add a checkbox for visibility
        visibility_checkbox = QCheckBox(self)
        visibility_checkbox.setChecked(True)  # Checked by default
        visibility_checkbox.stateChanged.connect(self.update_signal_visibility)
        table.setCellWidget(row, 3, visibility_checkbox)

    def update_signal_visibility(self):
        for row, (_, graph_number) in enumerate(self.imported_files):
            checkbox = None
            plot_item = None
            if graph_number == 0:
                checkbox = self.signals_info_table_1.cellWidget(row, 3)
                plot_item = self.plot_widget_1.getPlotItem()
            elif graph_number == 1:
                checkbox = self.signals_info_table_2.cellWidget(row, 3)
                plot_item = self.plot_widget_2.getPlotItem()

            if checkbox is not None and plot_item is not None:
                data_items = plot_item.listDataItems()
                if row < len(data_items):
                    if checkbox.isChecked():
                        # If the checkbox is checked, resume the timer and set the pen color
                        if (
                            self.number_of_signals_in_graph(graph_number=graph_number)
                            == 1
                        ):
                            self.graph_map[graph_number]["timer"].start(60)
                        data_items[row].setPen("black")
                    else:
                        # If the checkbox is unchecked, pause the timer and hide the signal
                        if (
                            self.number_of_signals_in_graph(graph_number=graph_number)
                            == 1
                        ):
                            self.graph_map[graph_number]["timer"].stop()
                        data_items[row].setPen(None)

    def number_of_signals_in_graph(self, graph_number):
        count = 0
        for _, signal_graph_number in self.imported_signals:
            if signal_graph_number == graph_number:
                count += 1
        return count

    def update_color_signal_list(self):
        for row, (_, graph_number) in enumerate(self.imported_files):
            color = QColor(
                self.signal_colors[row]
            )  # Convert the color string to QColor
            brush = QBrush(color)  # Create a QBrush with the color

            if graph_number == 0:
                color_label = QTableWidgetItem()
                color_label.setBackground(
                    brush
                )  # Set the background color using QBrush
                self.signals_info_table_1.setItem(row, 2, color_label)
            elif graph_number == 1:
                color_label = QTableWidgetItem()
                color_label.setBackground(
                    brush
                )  # Set the background color using QBrush
                self.signals_info_table_2.setItem(row, 2, color_label)

    def change_signal_color(self, current_row):
        if current_row >= 0:
            # Get the graph_number of the selected signal
            _, graph_number = self.imported_files[current_row]

            color = QColorDialog.getColor()
            if color.isValid():
                self.signal_colors[current_row] = color.name()

                # Update the corresponding signal color in the right graph's info table
                if graph_number == 0:
                    self.signals_info_table_1.item(current_row, 2).setBackground(
                        QBrush(color)
                    )
                elif graph_number == 1:
                    self.signals_info_table_2.item(current_row, 2).setBackground(
                        QBrush(color)
                    )

                selected_graph = self.graph_selector.currentIndex()
                selected_plot = self.graph_map[selected_graph]["plot"]
                selected_plot.clear()  # Clear the existing plot

    def ensure_signal_colors_length(self, num_signals):
        if len(self.signal_colors) < num_signals:
            # Initialize with some default colors
            default_colors = [
                "#FF0000",
                "#00FF00",
                "#0000FF",
                "#FFFF00",
                "#FF00FF",
                "#00FFFF",
            ]
            self.signal_colors.extend(
                default_colors[len(self.signal_colors) : num_signals]
            )

    def get_random_signal_color(self, _):
        # Generate a random color in the format '#RRGGBB'
        color = "#{:02X}{:02X}{:02X}".format(
            random.randint(0, 255), random.randint(0, 255), random.randint(0, 255)
        )
        return color

    def switch_graph(self, selected_row):
        self.plot_widget_1.clear()
        self.plot_widget_2.clear()
        self.timer_interval = 60
        # Get the signal data and graph number from the selected row
        signal_data, graph_number = self.imported_signals[selected_row]
        file_path, graph_number = self.imported_files[selected_row]

        self.current_graph = graph_number

        # Toggle the graph_number between 0 and 1
        graph_number = 1 - graph_number

        # Update the imported_signals and imported_files lists for the new graph
        self.imported_signals[selected_row] = (signal_data, graph_number)
        self.imported_files[selected_row] = (file_path, graph_number)

        if graph_number == 0:
            # If the signal is displayed in Graph 1, update signal_data_1 with new data
            self.signal_data_1 = signal_data
            if len(self.signal_data_1) > 0:
                self.playing_state[0] = True
                self.timer_1.start(self.timer_interval)
        elif graph_number == 1:
            # If the signal is displayed in Graph 2, update signal_data_2 with new data
            self.signal_data_2 = signal_data
            if len(self.signal_data_2) > 0:
                self.playing_state[1] = True
                self.timer_2.start(self.timer_interval)
        # Update the play/pause button icon based on the selected graph's playing state
        self.update_play_pause_button_icon(self.playing_state[self.current_graph])
        self.update_plot_1()  # Update Graph 1
        self.update_plot_2()  # Update Graph 2

        # Determine the source and destination tables
        source_table = (
            self.signals_info_table_1
            if graph_number == 1
            else self.signals_info_table_2
        )
        dest_table = (
            self.signals_info_table_2
            if graph_number == 1
            else self.signals_info_table_1
        )

        # Move the signal row from the source table to the destination table
        source_table.removeRow(selected_row)
        dest_table.insertRow(dest_table.rowCount())
        for col in range(source_table.columnCount()):
            item = source_table.takeItem(selected_row, col)
            dest_table.setItem(dest_table.rowCount() - 1, col, item)

        # Update the signal list (if it's a custom function you've implemented)
        self.update_signal_list()

    def create_context_menu(self, position):
        context_menu = QMenu(self)

        # Determine which table the right-click occurred in
        if self.signals_info_table_1.underMouse():
            table = self.signals_info_table_1
        elif self.signals_info_table_2.underMouse():
            table = self.signals_info_table_2
        else:
            return  # No valid table found

        # Get the selected row index
        selected_row = table.selectionModel().currentIndex().row()
        current_row = table.currentRow()

        # Add "Switch Graph" action and connect it to the switch_graph function with the selected row
        switch_graph_action = context_menu.addAction("Switch Graph")
        switch_graph_action.triggered.connect(lambda: self.switch_graph(selected_row))

        change_color_action = context_menu.addAction("Change Color")
        change_color_action.triggered.connect(
            lambda: self.change_signal_color(current_row)
        )

        # Get the global position of the cursor
        cursor_position = QCursor.pos()

        # Offset the context menu to the middle of the screen
        offset = QPoint(-int(context_menu.width() / 2), -int(context_menu.height() / 2))
        context_menu.exec_(cursor_position + offset)

    def contextMenuEvent(self, event):
        self.create_context_menu(event.pos())

    # --- Buttons Methods --- #
    # ----------------------- #

    def update_play_pause_button_icon(self, is_playing):
        if is_playing:
            self.pause_play_button.setIcon(QIcon("Resources/Icons/pause_button.png"))
        else:
            self.pause_play_button.setIcon(QIcon("Resources/Icons/play_button.png"))
        pass

    def toggle_play_pause(self, graph_index, checked):
        selected_timer = self.graph_map[graph_index]["timer"]
        if checked:
            self.playing_state[graph_index] = False
            selected_timer.stop()
        else:
            self.playing_state[graph_index] = True
            selected_timer.start(60)

        # Update the play/pause button icon based on the graph's playing state
        self.update_play_pause_button_icon(self.playing_state[graph_index])

    def reset_signal_for_graph(self, graph_index):
        selected_timer = self.graph_map[graph_index]["timer"]
        selected_timer.stop()
        selected_plot = self.graph_map[graph_index]["plot"]
        selected_plot.setData([])
        if graph_index == 0:
            self.signal_index_1 = 0  # Reset the signal index for graph 1
        elif graph_index == 1:
            self.signal_index_2 = 0  # Reset the signal index for graph 2
        selected_timer.start(60)  # Start the timer for the selected graph

    def pause_play_toggle_event(self, checked):
        if self.linked_graphs:
            for graph_index in self.graph_map:
                self.toggle_play_pause(graph_index, checked)
        else:
            selected_graph = self.graph_selector.currentIndex()
            self.toggle_play_pause(selected_graph, checked)

    def reset_signal(self):
        if self.linked_graphs:
            for graph_index in self.graph_map:
                self.reset_signal_for_graph(graph_index)
        else:
            selected_graph = self.graph_selector.currentIndex()
            self.reset_signal_for_graph(selected_graph)

    def apply_zoom(self, plot_widget, zoom_factor):
        x_range, y_range = plot_widget.plotItem.getViewBox().viewRange()
        center_x, center_y = sum(x_range) / 2, sum(y_range) / 2
        new_x_range = [
            (x_range[0] - center_x) * zoom_factor + center_x,
            (x_range[1] - center_x) * zoom_factor + center_x,
        ]
        new_y_range = [
            (y_range[0] - center_y) * zoom_factor + center_y,
            (y_range[1] - center_y) * zoom_factor + center_y,
        ]

        # Apply Y-axis zoom limit
        plot_widget.plotItem.getViewBox().setRange(
            xRange=new_x_range, yRange=new_y_range
        )

    def zoom_in_event(self):
        zoom_factor = 0.7  # Zoom factor (change as needed)

        if self.linked_graphs:
            for graph_index in range(2):  # Assuming there are two graphs
                selected_plot = (
                    self.plot_widget_1 if graph_index == 0 else self.plot_widget_2
                )
                self.apply_zoom(selected_plot, zoom_factor)
        else:
            selected_graph = self.graph_selector.currentIndex()
            selected_plot = (
                self.plot_widget_1 if selected_graph == 0 else self.plot_widget_2
            )
            self.apply_zoom(selected_plot, zoom_factor)

    def zoom_out_event(self):
        zoom_factor = 1.1  # Zoom factor (change as needed)

        if self.linked_graphs:
            for graph_index in range(2):  # Assuming there are two graphs
                selected_plot = (
                    self.plot_widget_1 if graph_index == 0 else self.plot_widget_2
                )
                self.apply_zoom(selected_plot, zoom_factor)
        else:
            selected_graph = self.graph_selector.currentIndex()
            selected_plot = (
                self.plot_widget_1 if selected_graph == 0 else self.plot_widget_2
            )
            self.apply_zoom(selected_plot, zoom_factor)

    def link_graphs_changed(self, state):
        if state == Qt.Checked:
            self.linked_graphs = True
        else:
            self.linked_graphs = False

        self.reset_signal()
        pass

    def update_speed(self):
        if self.linked_graphs:
            speed_value = self.speed_slider.value()
            self.speed_label.setText(f"Speed: x{speed_value}")
            for graph_index in self.graph_map:
                selected_timer = self.graph_map[graph_index]["timer"]
                interval = max(1, 60 - speed_value)
                selected_timer.setInterval(interval)
        else:
            selected_graph = self.graph_selector.currentIndex()
            selected_timer = self.graph_map[selected_graph]["timer"]
            speed_value = self.speed_slider.value()
            self.speed_label.setText(f"Speed: x{speed_value}")
            interval = max(1, 60 - speed_value)
            selected_timer.setInterval(interval)
        pass

    def graph_selected(self, index):
        # Store the last position of the currently selected graph
        selected_graph = (
            self.graph_selector.currentIndex()
        )  # returns the index of the currently selected item to remember it when changed
        if selected_graph == 0:
            self.last_position[0] = self.signal_index_1
        elif selected_graph == 1:
            self.last_position[1] = self.signal_index_2

        # Update the selected graph based on the dropdown list
        selected_graph = index
        selected_timer = self.graph_map[selected_graph]["timer"]
        selected_timer.stop()

        if selected_graph == 0:
            self.signal_index_1 = self.last_position[0]  # Set to the last position
            self.current_graph = 0
        elif selected_graph == 1:
            self.signal_index_2 = self.last_position[1]  # Set to the last position
            self.current_graph = 1

        selected_plot = self.graph_map[selected_graph]["plot"]
        selected_plot.setData([])
        selected_timer.start(60)  # Start the timer for the selected graph

        # Update the play/pause button icon based on the selected graph's playing state
        self.update_play_pause_button_icon(self.playing_state[self.current_graph])
        pass

    def take_snapshot_event(self):
        # Get the selected graph number from the graph_selector
        selected_graph = (
            self.graph_selector.currentIndex()
        )  # assuming graph_selector is a QComboBox

        # Find all the signal data for the selected graph number
        selected_signals = []

        for _, graph_number in self.imported_files:
            if graph_number == selected_graph:
                if selected_graph == 0:
                    selected_signals.append(self.signal_data_1[: self.signal_index_1])
                elif selected_graph == 1:
                    selected_signals.append(self.signal_data_2[: self.signal_index_2])

        if selected_signals:
            # Get or create a document file for snapshots
            snapshot_doc_path = "signal_snapshots.docx"
            if os.path.exists(snapshot_doc_path):
                doc = Document(snapshot_doc_path)
            else:
                doc = Document()
                doc.add_heading("Signal Snapshots", level=1)

            # Prompt the user to choose a directory for saving the file
            options = QFileDialog.Options()
            options |= QFileDialog.ReadOnly
            file_dialog = QFileDialog(self)
            file_dialog.setFileMode(QFileDialog.DirectoryOnly)
            selected_dir = file_dialog.getExistingDirectory(
                self, "Select Directory", "", options=options
            )

            if selected_dir:
                # Create a new document or open an existing one
                file_path = os.path.join(selected_dir, "signal_snapshot.docx")
                if not os.path.exists(file_path):
                    doc = Document()
                    doc.add_heading("Signal Snapshots", level=1)
                else:
                    doc = Document(file_path)

                # Export the plot_widget_1 or plot_widget_2 scene to a PNG file
                if selected_graph == 0:
                    exporter = pg.exporters.ImageExporter(
                        self.plot_widget_1.getPlotItem()
                    )
                elif selected_graph == 1:
                    exporter = pg.exporters.ImageExporter(
                        self.plot_widget_2.getPlotItem()
                    )

                temp_file_path = "temp_snapshot.png"
                exporter.export(temp_file_path)
                doc.add_picture(temp_file_path, width=Inches(6))
                os.remove(temp_file_path)

                # Add signal statistics to the document
                doc.add_heading("Signal Statistics", level=1)
                for i, signal_data in enumerate(selected_signals):
                    signal_stats = self.calculate_signal_stats(signal_data)
                    doc.add_paragraph(f"Signal {i + 1} Statistics:")
                    for key, value in signal_stats.items():
                        doc.add_paragraph(f"{key}: {value:.2f}")

                # Save the document to the selected directory
                doc.save(file_path)
                print(f"Snapshot added to the document and saved to: {file_path}")
        else:
            print(f"No signals found for graph {selected_graph}.")

    def convert_to_pdf(self):
        # Prompt the user to choose an existing DOCX file
        options = QFileDialog.Options()
        options |= QFileDialog.ReadOnly
        file_dialog = QFileDialog(self)
        selected_file, _ = file_dialog.getOpenFileName(
            self,
            "Select DOCX File to Convert to PDF",
            "",
            "DOCX Files (*.docx)",
            options=options,
        )

        if selected_file:
            # Generate PDF from the selected DOCX file
            pdf_file = selected_file.replace(".docx", ".pdf")
            convert(selected_file, pdf_file)
            print(f"PDF report generated successfully: {pdf_file}")
        else:
            print("PDF generation canceled. Please select a DOCX file to convert.")

    def calculate_signal_stats(self, signal_data):
        selected_graph = self.graph_selector.currentIndex()
        if signal_data is not None:
            mean_value = np.mean(signal_data)
            std_deviation = np.std(signal_data)
            duration = len(
                signal_data
            )  # Assuming the duration is the length of the signal
            min_value = np.min(signal_data)
            max_value = np.max(signal_data)
            return {
                "Mean": mean_value,
                "Standard Deviation": std_deviation,
                "Duration": duration,
                "Min Value": min_value,
                "Max Value": max_value,
            }

    def get_signal_stats(self, selected_graph):
        if selected_graph == 0:
            signal_data = self.signal_data_1
        elif selected_graph == 1:
            signal_data = self.signal_data_2
        return self.calculate_signal_stats(signal_data)

    def openDocumentation(self):
        # Open the specified URL in the default web browser
        webbrowser.open("https://github.com/cln-Kafka/SignalViewer/tree/Task-1")


if __name__ == "__main__":
    app = QApplication(sys.argv)

    with open("Diffnes.qss", "r") as f:
        stylesheet = f.read()
        app.setStyleSheet(stylesheet)

    window = SignalViewer()
    window.show()
    sys.exit(app.exec_())
